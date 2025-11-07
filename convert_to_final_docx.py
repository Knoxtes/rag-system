#!/usr/bin/env python3
"""
Convert final production analysis to Word document format
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Blue
    return heading

def add_table_with_style(doc, rows, cols, style='Light Grid Accent 1'):
    """Create a table with styling"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    return table

def create_word_document():
    """Create comprehensive Word document from final analysis"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Production Cost Analysis', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph('350 Users RAG System - Complete Comparison')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # System Parameters
    add_heading(doc, 'System Parameters', 1)
    params = [
        'Users: 350 active users',
        'Data Corpus: 4,529 text files, 7.94 GB, 2.132 billion tokens',
        'Images: 1,096 portfolio images',
        'Base Usage: 20 requests per user per day',
        'Cache Hit Rate: 15% (conservative estimate)',
        'RAG Refresh: Weekly (4 times per month)'
    ]
    for param in params:
        doc.add_paragraph(param, style='List Bullet')
    
    doc.add_paragraph()
    
    # Token Calculations
    add_heading(doc, 'Token Calculations', 1)
    
    add_heading(doc, 'Per Request Flow', 2)
    doc.add_paragraph('1. User Query: 50 tokens')
    doc.add_paragraph('2. Vector Search: Retrieves 3 most relevant documents')
    doc.add_paragraph('3. Context Window: 3 documents × 1,500 tokens = 4,500 tokens')
    doc.add_paragraph('4. LLM Response: 400 tokens generated')
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Total per request:\n').bold = True
    doc.add_paragraph('• Input to LLM: 4,500 tokens', style='List Bullet')
    doc.add_paragraph('• Output from LLM: 400 tokens', style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading(doc, 'Monthly Request Volume', 2)
    para = doc.add_paragraph()
    para.add_run('Total Requests: ').bold = True
    para.add_run('350 users × 20 req/day × 30 days = 210,000 requests/month')
    
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('With 15% Cache Hit Rate:\n').bold = True
    doc.add_paragraph('• Cached responses (free): 31,500 requests', style='List Bullet')
    doc.add_paragraph('• API calls to Gemini: 178,500 requests/month', style='List Bullet')
    
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Token Usage (API calls only):\n').bold = True
    doc.add_paragraph('• Input tokens: 803,250,000 tokens/month (803.25M)', style='List Bullet')
    doc.add_paragraph('• Output tokens: 71,400,000 tokens/month (71.4M)', style='List Bullet')
    
    doc.add_page_break()
    
    # STACK 1
    add_heading(doc, 'STACK 1: Full Local Processing 🏠', 1)
    
    doc.add_paragraph('Architecture: Local embeddings, local OCR, local vector DB, Gemini Flash API only')
    doc.add_paragraph()
    
    # Components table
    add_heading(doc, 'Components', 2)
    table = add_table_with_style(doc, 6, 4)
    headers = ['Component', 'Solution', 'Location', 'Cost']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['LLM', 'Gemini 2.0 Flash', 'Google API', '$81.66/month'],
        ['Embeddings', 'sentence-transformers', 'Local server', '$0'],
        ['Vector DB', 'ChromaDB (SQLite)', 'Local server', '$0'],
        ['OCR', 'EasyOCR', 'Local server', '$0'],
        ['Cache', 'In-memory dict', 'Local server', '$0'],
    ]
    
    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_paragraph()
    
    # Cost breakdown
    add_heading(doc, 'Cost Breakdown', 2)
    
    add_heading(doc, '1. Gemini Flash API (User Queries)', 3)
    para = doc.add_paragraph()
    para.add_run('Pricing:\n').bold = True
    doc.add_paragraph('• Input: $0.075 per 1M tokens', style='List Bullet')
    doc.add_paragraph('• Output: $0.30 per 1M tokens', style='List Bullet')
    
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Monthly Cost:\n').bold = True
    doc.add_paragraph('• Input: 803.25M × $0.075/1M = $60.24', style='List Bullet')
    doc.add_paragraph('• Output: 71.4M × $0.30/1M = $21.42', style='List Bullet')
    doc.add_paragraph('• Total: $81.66/month', style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading(doc, '2. Embedding Generation (Weekly RAG Digest)', 3)
    doc.add_paragraph('Initial Digest (one-time): 47 minutes, $0')
    doc.add_paragraph('Weekly Re-digest: 4.7 minutes, $0')
    doc.add_paragraph('Monthly time: 19 minutes/month, $0')
    
    doc.add_paragraph()
    
    add_heading(doc, '3. OCR Processing (Weekly)', 3)
    doc.add_paragraph('Initial Processing (one-time): 64 minutes, $0')
    doc.add_paragraph('Weekly Re-processing: 64 minutes per week')
    doc.add_paragraph('Monthly time: 256 minutes = 4.3 hours/month, $0')
    
    doc.add_paragraph()
    
    add_heading(doc, '4. Vector Database Storage', 3)
    doc.add_paragraph('Storage: 3.27 GB (SQLite)')
    doc.add_paragraph('Cost: $0 (local storage)')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Total Monthly Cost: ').bold = True
    run = para.add_run('$81.66')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_page_break()
    
    # Server Requirements Stack 1
    add_heading(doc, 'STACK 1: Server Resource Requirements', 1)
    
    add_heading(doc, 'Hardware Specifications', 2)
    para = doc.add_paragraph()
    para.add_run('Minimum Server:\n').bold = True
    doc.add_paragraph('• CPU: 8 vCPU', style='List Bullet')
    doc.add_paragraph('• RAM: 16 GB', style='List Bullet')
    doc.add_paragraph('• Storage: 20 GB', style='List Bullet')
    doc.add_paragraph('• Network: Standard', style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading(doc, 'CPU Usage Breakdown', 2)
    doc.add_paragraph('Idle State: 5-10%')
    doc.add_paragraph('Per Request: 350ms at ~20% CPU spike')
    doc.add_paragraph('Peak Load (70 concurrent users): 40-60% sustained')
    doc.add_paragraph('Weekly RAG Digest: 80-90% for 4.7 minutes (4 times/month)')
    doc.add_paragraph('Weekly OCR: 70-85% for 64 minutes (4 times/month)')
    
    doc.add_paragraph()
    
    add_heading(doc, 'RAM Usage Breakdown', 2)
    doc.add_paragraph('Base Memory: 2.2 GB')
    doc.add_paragraph('Peak Concurrent (70 users): 3.5 GB additional')
    doc.add_paragraph('Total Peak RAM: 5.7 GB (fits in 16 GB)')
    
    doc.add_paragraph()
    
    add_heading(doc, 'Storage Breakdown', 2)
    table = add_table_with_style(doc, 8, 3)
    headers = ['Item', 'Size', 'Type']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['Text corpus', '7.94 GB', 'Raw files'],
        ['Vector database', '3.27 GB', 'SQLite'],
        ['Models (cached)', '1.6 GB', 'PyTorch'],
        ['Application code', '100 MB', 'Python'],
        ['Logs (monthly)', '300 MB', 'Text'],
        ['OS overhead', '4 GB', 'System'],
        ['Total', '17.2 GB', '20 GB sufficient'],
    ]
    
    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_page_break()
    
    # STACK 2
    add_heading(doc, 'STACK 2: Full Google Cloud Solutions ☁️', 1)
    
    doc.add_paragraph('Architecture: Vertex AI embeddings, Google Vision OCR, ChromaDB local, Gemini Flash API')
    doc.add_paragraph()
    
    # Components table
    add_heading(doc, 'Components', 2)
    table = add_table_with_style(doc, 6, 4)
    headers = ['Component', 'Solution', 'Location', 'Cost']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['LLM', 'Gemini 2.0 Flash', 'Google API', '$81.66/month'],
        ['Embeddings', 'Vertex AI text-embedding-004', 'Google API', '$0.10/month'],
        ['Vector DB', 'ChromaDB (SQLite)', 'Local server', '$0'],
        ['OCR', 'Google Cloud Vision', 'Google API', '$5.08/month'],
        ['Cache', 'In-memory dict', 'Local server', '$0'],
    ]
    
    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_paragraph()
    
    # Cost breakdown Stack 2
    add_heading(doc, 'Cost Breakdown', 2)
    doc.add_paragraph('1. Gemini Flash API: $81.66/month (same as Stack 1)')
    doc.add_paragraph('2. Vertex AI Embeddings: $0.10/month (3.93M tokens)')
    doc.add_paragraph('3. Google Cloud Vision OCR: $5.08/month (3,384 billable images)')
    doc.add_paragraph('4. Vector Database: $0 (local storage)')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Total Monthly Cost: ').bold = True
    run = para.add_run('$86.84')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('One-Time Setup Costs:\n').bold = True
    doc.add_paragraph('• Initial Vertex AI Digest: $53.30', style='List Bullet')
    doc.add_paragraph('• Initial Google Vision OCR: $0.14', style='List Bullet')
    doc.add_paragraph('• Total One-Time: $53.44', style='List Bullet')
    
    doc.add_page_break()
    
    # Server Requirements Stack 2
    add_heading(doc, 'STACK 2: Server Resource Requirements', 1)
    
    add_heading(doc, 'Hardware Specifications', 2)
    para = doc.add_paragraph()
    para.add_run('Minimum Server (half the requirement vs Stack 1):\n').bold = True
    doc.add_paragraph('• CPU: 4 vCPU', style='List Bullet')
    doc.add_paragraph('• RAM: 4 GB', style='List Bullet')
    doc.add_paragraph('• Storage: 15 GB', style='List Bullet')
    doc.add_paragraph('• Network: Standard', style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading(doc, 'CPU Usage Breakdown', 2)
    doc.add_paragraph('Idle State: 5-10%')
    doc.add_paragraph('Per Request: 100ms at ~6% CPU spike')
    doc.add_paragraph('Peak Load (70 concurrent users): 10-20% sustained')
    doc.add_paragraph('Weekly RAG Digest: 10-15% for 1 minute')
    doc.add_paragraph('Weekly OCR: 5-10% for 5 minutes')
    
    doc.add_paragraph()
    
    add_heading(doc, 'RAM Usage Breakdown', 2)
    doc.add_paragraph('Base Memory: 700 MB')
    doc.add_paragraph('Peak Concurrent (70 users): 700 MB additional')
    doc.add_paragraph('Total Peak RAM: 1.4 GB (fits in 4 GB easily)')
    
    doc.add_page_break()
    
    # Comparison Summary
    add_heading(doc, 'Stack Comparison Summary', 1)
    
    table = add_table_with_style(doc, 18, 3)
    headers = ['Metric', 'Stack 1: Full Local', 'Stack 2: Full Google']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['Monthly Cost', '$81.66', '$86.84'],
        ['One-time Setup Cost', '$0', '$53.44'],
        ['Server CPU Required', '8 vCPU', '4 vCPU'],
        ['Server RAM Required', '16 GB', '4 GB'],
        ['Storage Required', '20 GB', '15 GB'],
        ['Idle CPU Usage', '5-10%', '5-10%'],
        ['Peak CPU Usage', '40-60% sustained', '10-20% sustained'],
        ['Peak RAM Usage', '5.7 GB', '1.4 GB'],
        ['Request Latency', '350ms', '100ms'],
        ['Weekly Maintenance Time', '68.7 min high CPU', '6 min low CPU'],
        ['Initial Setup Time', '111 minutes', '15 minutes'],
        ['Privacy', '✅ Fully local', '⚠️ Data to Google'],
        ['External Dependencies', 'Gemini API only', '3 Google APIs'],
        ['Server Strain', 'Medium-High', 'Low'],
        ['Annual Cost', '$979.92', '$1,042.08'],
        ['Annual Savings', '$62.16 cheaper', '-'],
        ['Performance', 'Standard', '3.5× faster'],
    ]
    
    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_page_break()
    
    # Scaling tables
    add_heading(doc, 'Scaling: Requests Per Day (350 Users)', 1)
    
    add_heading(doc, 'Stack 1: Full Local Costs', 2)
    table = add_table_with_style(doc, 9, 6)
    headers = ['Req/Day', 'Monthly Reqs', 'API Calls', 'Input Tokens', 'Output Tokens', 'Monthly Cost']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['5', '52,500', '44,625', '200.8M', '17.85M', '$20.41'],
        ['10', '105,000', '89,250', '401.6M', '35.7M', '$40.83'],
        ['15', '157,500', '133,875', '602.4M', '53.55M', '$61.24'],
        ['20', '210,000', '178,500', '803.25M', '71.4M', '$81.66'],
        ['25', '262,500', '223,125', '1,004M', '89.25M', '$102.07'],
        ['30', '315,000', '267,750', '1,205M', '107.1M', '$122.49'],
        ['40', '420,000', '357,000', '1,607M', '142.8M', '$163.32'],
        ['50', '525,000', '446,250', '2,008M', '178.5M', '$204.15'],
    ]
    
    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_paragraph()
    
    add_heading(doc, 'Stack 2: Full Google Costs', 2)
    table = add_table_with_style(doc, 9, 5)
    headers = ['Req/Day', 'Gemini Flash', 'Vertex AI', 'Google Vision', 'Total Cost']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['5', '$20.41', '$0.05', '$5.08', '$25.54'],
        ['10', '$40.83', '$0.08', '$5.08', '$45.99'],
        ['15', '$61.24', '$0.09', '$5.08', '$66.41'],
        ['20', '$81.66', '$0.10', '$5.08', '$86.84'],
        ['25', '$102.07', '$0.11', '$5.08', '$107.26'],
        ['30', '$122.49', '$0.12', '$5.08', '$127.69'],
        ['40', '$163.32', '$0.14', '$5.08', '$168.54'],
        ['50', '$204.15', '$0.16', '$5.08', '$209.39'],
    ]
    
    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_page_break()
    
    # User growth scaling
    add_heading(doc, 'Scaling: User Growth (20 Req/Day)', 1)
    
    table = add_table_with_style(doc, 8, 6)
    headers = ['Users', 'Monthly Reqs', 'Stack 1 Cost', 'Stack 2 Cost', 'Stack 1 CPU', 'Stack 2 CPU']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ['100', '60,000', '$23.33', '$24.81', '20-30%', '5-10%'],
        ['250', '150,000', '$58.33', '$62.04', '35-50%', '10-15%'],
        ['350', '210,000', '$81.66', '$86.84', '40-60%', '10-20%'],
        ['500', '300,000', '$116.66', '$124.06', '55-75%', '15-25%'],
        ['750', '450,000', '$175.00', '$186.09', '70-90%', '20-35%'],
        ['1,000', '600,000', '$233.33', '$248.12', 'Upgrade needed', '25-40%'],
        ['2,000', '1,200,000', '$466.66', '$496.24', 'Cluster needed', '45-65%'],
    ]
    
    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_page_break()
    
    # Recommendations
    add_heading(doc, 'Recommendations', 1)
    
    add_heading(doc, 'Choose Stack 1 (Full Local) if:', 2)
    reasons = [
        '✅ Budget is critical ($62/year savings)',
        '✅ Data privacy is paramount (embeddings/OCR stay local)',
        '✅ You have server capacity (16GB RAM, 8 vCPU available)',
        '✅ 350ms latency is acceptable',
        '✅ Willing to manage weekly maintenance (68 minutes high CPU)',
        '✅ Growing beyond 350 users is unlikely'
    ]
    for reason in reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading(doc, 'Choose Stack 2 (Full Google) if:', 2)
    reasons = [
        '✅ Performance is critical (3.5× faster)',
        '✅ Server resources are limited (75% less RAM, 50% less CPU)',
        '✅ Low latency required (100ms response time)',
        '✅ Minimal maintenance preferred (6 min/week vs 68 min)',
        '✅ Planning to scale (handles 2× users on same hardware)',
        '✅ $5.18/month extra is acceptable'
    ]
    for reason in reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading(doc, 'Optimal Hybrid (Recommended) ⭐', 2)
    doc.add_paragraph('Use Stack 1 base with selective Google services:')
    doc.add_paragraph('• Gemini Flash: API (required)', style='List Bullet')
    doc.add_paragraph('• Embeddings: Local (save $0.10/mo + privacy)', style='List Bullet')
    doc.add_paragraph('• OCR: Google Vision (save 256 min/month for $5.08)', style='List Bullet')
    doc.add_paragraph('• Vector DB: Local (free)', style='List Bullet')
    
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Hybrid Cost: ').bold = True
    para.add_run('$86.74/month ($81.66 + $5.08)')
    
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Benefits: ').bold = True
    para.add_run('Privacy for embeddings, time savings on OCR, low cost')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_para.add_run('Document Version: 4.0 (Complete rebuild, two-stack comparison)\nLast Updated: November 6, 2025')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    return doc

if __name__ == '__main__':
    print("Creating Word document from final analysis...")
    doc = create_word_document()
    
    output_file = 'PRODUCTION_ANALYSIS_FINAL.docx'
    doc.save(output_file)
    print(f"✓ Word document saved: {output_file}")
    print(f"✓ Complete two-stack comparison with server strain analysis")
    print(f"✓ Includes scaling tables and detailed recommendations")
